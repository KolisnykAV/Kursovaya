import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement, ns
from docx.shared import Cm


class FileProcessorApp:
 def __init__(self, master):
     self.master = master
     self.master.title("File Processor App")

     self.button_style = ("Times New Roman", 12)

     self.select_button = tk.Button(self.master, text="Выбрать файл", font=self.button_style, command=self.select_file)
     self.select_button.pack(pady=10)

     self.process_button = tk.Button(self.master, text="Обработать файл", font=self.button_style, command=self.process_file)
     self.process_button.pack(pady=10)

 def create_element(self, name):
     return OxmlElement(name)

 def create_attribute(self, element, name, value):
     element.set(ns.qn(name), value)

 def add_page_number(self, run):
     fldChar1 = self.create_element('w:fldChar')
     self.create_attribute(fldChar1, 'w:fldCharType', 'begin')

     instrText = self.create_element('w:instrText')
     self.create_attribute(instrText, 'xml:space', 'preserve')
     instrText.text = "PAGE"

     fldChar2 = self.create_element('w:fldChar')
     self.create_attribute(fldChar2, 'w:fldCharType', 'end')

     run._r.append(fldChar1)
     run._r.append(instrText)
     run._r.append(fldChar2)

 def select_file(self):
     self.file_path = filedialog.askopenfilename(filetypes=[("Word documents", "*.docx"), ("All files", "*.*")])
     messagebox.showinfo("Выбран файл", f"Выбран файл: {self.file_path}")


 def process_file(self):

     if self.file_path:
         doc = Document(self.file_path)
         section = doc.sections[0]
         section.top_margin = Cm(2)
         section.bottom_margin = Cm(2)
         section.left_margin = Cm(2)
         section.right_margin = Cm(2)
         for paragraph in doc.paragraphs:
             if paragraph.style.name.startswith('Heading'):
                 heading_level = int(paragraph.style.name.split()[1])

                 if heading_level == 1:
                     paragraph_format = paragraph.paragraph_format
                     paragraph_format.space_after = Pt(12)
                     paragraph_format.space_before = Pt(12)
                     paragraph_format.left_indent = Pt(36)
                     paragraph_format.right_indent = Pt(36)
                     paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                     for run in paragraph.runs:
                         run.font.name = "Times New Roman"
                         run.font.size = Pt(16)
                         run.font.bold = True
                         run.font.all_caps = True
                 elif heading_level == 2:
                     paragraph_format = paragraph.paragraph_format
                     paragraph_format.space_after = Pt(12)
                     paragraph_format.space_before = Pt(12)
                     paragraph_format.left_indent = Pt(36)
                     paragraph_format.right_indent = Pt(36)
                     paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                     for run in paragraph.runs:
                         run.font.name = "Times New Roman"
                         run.font.size = Pt(15)
                         run.font.bold = True
                 elif heading_level == 3:
                     paragraph_format = paragraph.paragraph_format
                     paragraph_format.space_after = Pt(12)
                     paragraph_format.space_before = Pt(12)
                     paragraph_format.left_indent = Pt(36)
                     paragraph_format.right_indent = Pt(36)
                     paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                     for run in paragraph.runs:
                         run.font.name = "Times New Roman"
                         run.font.size = Pt(14)
                         run.font.bold = True
                         run.font.italic = True
             else:
                 for run in paragraph.runs:
                     run.font.name = "Times New Roman"
                     run.font.size = Pt(14)
                     paragraph_format = paragraph.paragraph_format
                     paragraph_format.first_line_indent = Cm(1)  # абзацный отступ 10 мм
                     # paragraph_format.line_spacing = 1.5   второй вариант межстрочного интервала
                     paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                     paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                     paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                     paragraph_format.space_after = Pt(0.3)  # в итоге интервал 1,2

         # добавить номер страницы
         footer = doc.sections[0].footer
         paragraph = footer.paragraphs[0]
         paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
         paragraph.paragraph_format.left_indent = Pt(0)
         paragraph.paragraph_format.right_indent = Pt(0)
         self.add_page_number(paragraph.add_run())

         file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word documents", "*.docx"), ("All files", "*.*")])

         if file_path:
             doc.save(file_path)
             print(f"Файл успешно сохранен по пути: {file_path}")
     else:
         messagebox.showwarning("Обработка файла", "Выберите файл перед обработкой.")

 def download_processed_file(self):
     directory_path = filedialog.askdirectory()
     if directory_path:
         print("Выбрана папка:", directory_path)

     # messagebox.showinfo("Скачивание файла", "Обработанный файл успешно скачан.")

if __name__ == "__main__":
  root = tk.Tk()
  app = FileProcessorApp(root)
  root.update()
  root.mainloop()
